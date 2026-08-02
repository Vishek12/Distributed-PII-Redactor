"""
DOCX Text Extractor Module
Handles text extraction from DOCX files via file path or raw bytes stream.
"""

import io 
from typing import Union, BinaryIO
import docx
from pathlib import Path

class DOCXExtractor:
    """Extracts raw text content from DOCX documents."""
    
    @staticmethod
    def extract_text(file_source: Union[str, BinaryIO, bytes]) -> str: 
    
        docx_stream = file_source
        if isinstance(file_source, bytes):
            docx_stream = io.BytesIO(file_source)

        try:
            doc = docx.Document(docx_stream)
        except Exception as e: 
            raise RuntimeError(f"Failed to extract text from DOCX: {str(e)}") from e

        extracted_elements = []


        #Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_elements.append(para.text.strip())


        #Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text: 
                    extracted_elements.append(" | ".join(row_text))

        #Extract text from headers and footers 
        for section in doc.sections:
            for header_para in section.header.paragraphs:
                if header_para.text.strip():
                    extracted_elements.append(header_para.text.strip())
            for footer_para in section.footer.paragraphs:
                if footer_para.text.strip():
                    extracted_elements.append(footer_para.text.strip()) 

        return "\n\n".join(extracted_elements)

    #Local testing of extracting DOCX text
if __name__ == "__main__":
    # Example usage / sanity check
    sample_docx_path = Path(__file__).resolve().parent.parent.parent / "data/sample.docx"
    print(DOCXExtractor.extract_text(sample_docx_path))