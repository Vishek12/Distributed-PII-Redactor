"""
Document Exporter Module
Converts redacted plain text or processed document streams into downloadable .docx, .pdf, or .txt byte streams.
"""

import io
import docx
import fitz  # PyMuPDF


class DocumentExporter: 

    @staticmethod
    def export_docx_bytes(text: str) -> bytes: 
        """Converts redacted text into an in-memory .docx file byte stream."""
        # Fix: Instantiate the Document object
        doc = docx.Document() 

        for paragraph in text.split("\n"): 
            if paragraph.strip():
                doc.add_paragraph(paragraph)
            else:
                doc.add_paragraph("")  # Preserve spacing for blank lines

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod 
    def export_text_bytes(text: str) -> bytes: 
        """Converts redacted text into an in-memory .txt file byte stream."""
        return text.encode("utf-8")

    @staticmethod
    def create_redacted_pdf_bytes(original_pdf_bytes: bytes, pii_entities: list[str]) -> bytes:
        """
        Searches for PII strings on the original PDF and applies visual black redaction boxes.
        
        :param original_pdf_bytes: Raw bytes of the uploaded PDF file.
        :param pii_entities: List of target PII strings to scrub (e.g., ["John Doe", "416-555-0199"]).
        :return: Redacted PDF byte stream.
        """
        doc = fitz.open(stream=original_pdf_bytes, filetype="pdf")

        # Deduplicate and filter out empty strings or single-character artifacts
        targets = set(entity.strip() for entity in pii_entities if entity and len(entity.strip()) > 1)

        for page in doc:
            for pii_text in targets:
                # 1. Locate coordinate bounding boxes (Rect objects) for the PII text
                text_instances = page.search_for(pii_text)

                # 2. Add redaction annotations over every match
                for rect in text_instances:
                    page.add_redact_annot(
                        rect,
                        fill=(0, 0, 0),        # Solid black fill box
                        text="[REDACTED]",    # Overlay text
                        text_color=(1, 1, 1), # White text
                        fontsize=8
                    )

            # 3. Permanently scrub underlying PDF stream data under annotations
            page.apply_redactions()

        output_buffer = io.BytesIO()
        doc.save(output_buffer, garbage=4, deflate=True)  # Optimize & compress
        doc.close()
        
        return output_buffer.getvalue()